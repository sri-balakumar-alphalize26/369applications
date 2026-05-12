"""
Generates Attendance_User_Manual.docx — a plain-English user manual for the
four attendance features (Office, Leave, Late Waiver, Field).

Run from the project root:
    python docs/_generate_manual.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm


OUT = Path(__file__).resolve().parent / "Attendance_User_Manual.docx"


# ---------- styling helpers ----------

NAVY = RGBColor(0x1F, 0x3A, 0x68)
ACCENT = RGBColor(0x21, 0x6E, 0x39)
TEXT = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x66, 0x66, 0x66)


def shade(cell, hex_colour):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=10, colour=TEXT, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = colour
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = "Calibri"
    if level == 0:
        run.font.size = Pt(28)
        run.font.color.rgb = NAVY
        run.bold = True
    elif level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = NAVY
        run.bold = True
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = NAVY
        run.bold = True
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = ACCENT
        run.bold = True
    return h


def add_para(doc, text, bold=False, italic=False, size=11, colour=TEXT, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = colour
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run("")
    p.runs[0].text = ""  # reset
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT
    return p


def add_callout(doc, label, body):
    """Coloured callout box for tips / warnings."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.cell(0, 0)
    shade(cell, "EAF3FF")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.color.rgb = NAVY
    r1.font.size = Pt(11)
    r1.font.name = "Calibri"
    r2 = p.add_run(body)
    r2.font.size = Pt(11)
    r2.font.color.rgb = TEXT
    r2.font.name = "Calibri"
    doc.add_paragraph()  # spacer
    return tbl


def add_rules_table(doc, rows):
    """Two-column 'Rule | Detail' table."""
    tbl = doc.add_table(rows=len(rows) + 1, cols=2)
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = True

    set_cell_text(tbl.cell(0, 0), "Rule", bold=True, size=11,
                  colour=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_text(tbl.cell(0, 1), "What it means", bold=True, size=11,
                  colour=RGBColor(0xFF, 0xFF, 0xFF))
    shade(tbl.cell(0, 0), "1F3A68")
    shade(tbl.cell(0, 1), "1F3A68")

    for i, (k, v) in enumerate(rows, start=1):
        set_cell_text(tbl.cell(i, 0), k, bold=True, size=10)
        set_cell_text(tbl.cell(i, 1), v, size=10)
    doc.add_paragraph()


def add_status_table(doc):
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Light Grid Accent 1"
    set_cell_text(tbl.cell(0, 0), "Colour", bold=True, size=11,
                  colour=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_text(tbl.cell(0, 1), "Status it means", bold=True, size=11,
                  colour=RGBColor(0xFF, 0xFF, 0xFF))
    shade(tbl.cell(0, 0), "1F3A68")
    shade(tbl.cell(0, 1), "1F3A68")

    rows = [
        ("Gray",   "9E9E9E", "Draft or Cancelled — request not yet acted on, or you withdrew it"),
        ("Orange", "FF9800", "Pending — waiting for your manager / HR to review"),
        ("Green",  "4CAF50", "Approved — request accepted"),
        ("Red",    "F44336", "Rejected — request turned down (you'll usually see a reason)"),
    ]
    for i, (name, hex_, desc) in enumerate(rows, start=1):
        set_cell_text(tbl.cell(i, 0), name, bold=True, size=10,
                      colour=RGBColor(0xFF, 0xFF, 0xFF))
        shade(tbl.cell(i, 0), hex_)
        set_cell_text(tbl.cell(i, 1), desc, size=10)
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


# ---------- document content ----------

def build():
    doc = Document()

    # Default font for the document
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Reduce default margins slightly so manual feels spacious but not airy
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ---------- COVER ----------
    cover_para = doc.add_paragraph()
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_para.add_run("\n\n\n")
    run.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Attendance App")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("User Manual")
    r.font.size = Pt(22)
    r.font.color.rgb = ACCENT
    r.font.name = "Calibri"

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = intro.add_run(
        "A simple, step-by-step guide for marking your attendance, applying "
        "for leave, requesting late forgiveness, and recording field visits."
    )
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = GREY
    r.font.name = "Calibri"

    doc.add_paragraph()
    doc.add_paragraph()

    # Table of contents (manual list, not auto-generated)
    add_heading(doc, "What's inside", level=2)
    add_bullet(doc, "1. Before you start — how to log in and what you'll see")
    add_bullet(doc, "2. Office Attendance — check in and out at the office")
    add_bullet(doc, "3. Leave Request — apply for a day (or more) off")
    add_bullet(doc, "4. Late Waiver Request — ask to forgive a late mark")
    add_bullet(doc, "5. Field Attendance — for customer visits and on-trip days")
    add_bullet(doc, "Appendix — what the status colours mean")

    page_break(doc)

    # ---------- SECTION 1: BEFORE YOU START ----------
    add_heading(doc, "1. Before you start", level=1)
    add_para(doc,
             "This section covers the basics every employee needs to know before using "
             "any attendance feature. The app works the same way whether you're at the "
             "office or on the road.")

    add_heading(doc, "Logging in", level=2)
    add_numbered(doc, "Open the Attendance app on your phone.")
    add_numbered(doc, "Enter your username and password (the same ones your HR gave you).")
    add_numbered(doc, "Tap Login. You'll land on the Home screen.")

    add_heading(doc, "Identity check — fingerprint or PIN", level=2)
    add_para(doc,
             "Before you can mark attendance or submit a request, the app asks you to "
             "prove it's really you. There are two ways:")
    add_bullet(doc, "Fingerprint (recommended) — just touch the sensor when the prompt appears.")
    add_bullet(doc, "4-digit PIN — type your PIN if your phone doesn't support fingerprint or it isn't working.")
    add_callout(doc, "Tip:",
                "If the fingerprint scan fails twice, tap \"Use device PIN\" to fall back.")

    add_heading(doc, "Finding what you need on the Home screen", level=2)
    add_para(doc,
             "From Home, tap Attendance. You'll see options for the four features covered "
             "in this manual. Each one is described in detail in the sections that follow.")

    page_break(doc)

    # ---------- SECTION 2: OFFICE ATTENDANCE ----------
    add_heading(doc, "2. Office Attendance", level=1)
    add_para(doc,
             "Use this every working day when you arrive at and leave the office.",
             italic=True, colour=GREY)

    add_heading(doc, "Step by step", level=2)
    add_numbered(doc, "Open the app and go to Attendance → Punching.")
    add_numbered(doc, "Tap today's date on the calendar.")
    add_numbered(doc, "Tap the Check In button.")
    add_numbered(doc, "Scan your fingerprint (or enter PIN) when prompted.")
    add_numbered(doc, "The app checks your location. You must be within 100 metres of the office.")
    add_numbered(doc, "The camera opens — your photo is taken automatically (3-2-1 countdown).")
    add_numbered(doc, "Done! You'll see a confirmation that you've checked in.")
    add_para(doc, "Repeat the same steps when you leave — the button will say Check Out.", italic=True)

    add_heading(doc, "Rules at a glance", level=2)
    add_rules_table(doc, [
        ("Location",         "You must be within 100 m of the registered office to check in."),
        ("Photo",            "A photo is taken automatically every time you check in or out."),
        ("Punches per day",  "Up to 4: Forenoon In, Forenoon Out, Afternoon In, Afternoon Out."),
        ("No double punch",  "Once you've checked out of a session, you can't check back into the same session that day."),
        ("Offline",          "If there's no internet, your check-in is saved on the phone and syncs later."),
    ])

    add_heading(doc, "If you arrive late", level=2)
    add_para(doc,
             "The app automatically detects late arrival. You'll see how many minutes "
             "late you are and the deduction amount (if any). A pop-up will ask you to "
             "type a reason — be honest and specific (e.g., \"traffic jam on highway\", "
             "\"medical emergency at home\").")
    add_callout(doc, "Good to know:",
                "The first 5 late arrivals each month are usually free (\"grace period\"). "
                "Your company may set a different number. After that, each late attracts a deduction.")

    add_heading(doc, "Common problems", level=2)
    add_bullet(doc,
               "\"You are 250 m away from office\" — move closer to the building and try again. "
               "If GPS is fuzzy, wait a few seconds and retry.")
    add_bullet(doc,
               "\"Camera won't open\" — make sure you granted the app camera permission in your phone settings.")
    add_bullet(doc,
               "\"I checked in but was actually late by mistake\" — submit a Late Waiver Request (see Section 4).")
    add_bullet(doc,
               "\"Saved offline. Will sync when online\" — this is normal. Open the app on Wi-Fi and it auto-syncs.")

    page_break(doc)

    # ---------- SECTION 3: LEAVE REQUEST ----------
    add_heading(doc, "3. Leave Request", level=1)
    add_para(doc,
             "Use this when you need a day off (or several). The request goes to your "
             "manager for approval.",
             italic=True, colour=GREY)

    add_heading(doc, "Step by step", level=2)
    add_numbered(doc, "Open the app and tap Attendance.")
    add_numbered(doc, "Pick Leave Request and verify yourself (fingerprint/PIN).")
    add_numbered(doc, "Make sure you're on the New Request tab.")
    add_numbered(doc, "Tap a Leave Type button (Sick, Casual, Annual, etc.).")
    add_numbered(doc, "Pick the From Date.")
    add_numbered(doc, "If it's only half a day, tick Half Day Leave. Otherwise pick a To Date.")
    add_numbered(doc, "Type the reason in the box. This is mandatory.")
    add_numbered(doc, "Tap Submit. Review the confirmation pop-up, then tap SUBMIT again.")
    add_para(doc, "You'll see a success message and the app jumps to the My Requests tab.", italic=True)

    add_heading(doc, "Leave types — when to use which", level=2)
    tbl = doc.add_table(rows=7, cols=2)
    tbl.style = "Light Grid Accent 1"
    set_cell_text(tbl.cell(0, 0), "Type", bold=True, colour=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_text(tbl.cell(0, 1), "Use it for", bold=True, colour=RGBColor(0xFF, 0xFF, 0xFF))
    shade(tbl.cell(0, 0), "1F3A68")
    shade(tbl.cell(0, 1), "1F3A68")
    leave_rows = [
        ("Sick Leave",      "You're unwell or visiting a doctor."),
        ("Casual Leave",    "Personal errands, family matters, short breaks."),
        ("Annual Leave",    "Planned vacation or holiday trips."),
        ("Personal Leave",  "Specific personal needs not covered by the others."),
        ("Emergency Leave", "Sudden, unavoidable situations (accident, bereavement)."),
        ("Other",           "Any reason not fitting the above — explain clearly in the reason box."),
    ]
    for i, (k, v) in enumerate(leave_rows, start=1):
        set_cell_text(tbl.cell(i, 0), k, bold=True, size=10)
        set_cell_text(tbl.cell(i, 1), v, size=10)
    doc.add_paragraph()

    add_heading(doc, "Rules at a glance", level=2)
    add_rules_table(doc, [
        ("Paid quota (default)", "12 paid days per year, up to 1 per month. Your company may set different limits."),
        ("Reason",               "Always required."),
        ("Dates",                "To Date must be on or after From Date. You can pick today or any future date."),
        ("Half day",             "Counts as 0.5 days — tick the Half Day Leave checkbox."),
        ("Unpaid days",          "If you exceed your paid quota, the extra days may be deducted from your salary."),
    ])

    add_heading(doc, "What happens after you submit", level=2)
    add_bullet(doc, "The status starts as Pending (orange).")
    add_bullet(doc, "Your manager reviews and either approves (green) or rejects (red).")
    add_bullet(doc, "If approved, you'll see the manager's name on the card. If rejected, you'll see the reason.")
    add_bullet(doc, "You can open the My Requests tab any time to check the status.")

    add_heading(doc, "Cancelling a request", level=2)
    add_para(doc,
             "If you change your mind, open the request on the My Requests tab and tap "
             "Cancel Request. You can cancel from Draft, Pending, or even Approved status. "
             "Once cancelled, the leave is no longer counted.")

    add_heading(doc, "Common problems", level=2)
    add_bullet(doc,
               "\"I picked wrong dates\" — cancel the request and create a fresh one.")
    add_bullet(doc,
               "\"Manager rejected — can I try again?\" — yes, submit a new request with a clearer reason.")
    add_bullet(doc,
               "\"Where do I see my leave balance?\" — your HR portal shows the official balance; the app shows your past requests on the My Requests tab.")

    page_break(doc)

    # ---------- SECTION 4: LATE WAIVER REQUEST ----------
    add_heading(doc, "4. Late Waiver Request", level=1)
    add_para(doc,
             "Use this when you were late for a valid reason and want the deduction "
             "forgiven (for example, you were on an office errand or had a genuine "
             "emergency).",
             italic=True, colour=GREY)

    add_heading(doc, "Step by step", level=2)
    add_numbered(doc, "Open the app and tap Attendance.")
    add_numbered(doc, "Pick Late Waiver Request and verify yourself.")
    add_numbered(doc, "Stay on the New Request tab.")
    add_numbered(doc,
                 "Scroll the list of your recent late records. Each shows the date, "
                 "minutes late, and deduction amount.")
    add_numbered(doc,
                 "Tap the circle next to the late record you want to waive. "
                 "(Records already waived are greyed out — you can't pick them.)")
    add_numbered(doc,
                 "Type the reason for waiver in the box. Be specific — e.g., \"client "
                 "meeting at customer site\", \"traffic block on main road due to roadwork\".")
    add_numbered(doc, "Tap Submit and confirm in the pop-up.")

    add_heading(doc, "Rules at a glance", level=2)
    add_rules_table(doc, [
        ("Time window",   "Only your late records from the last 30 days are eligible."),
        ("First late only", "Only the first late check-in of a day per session can be waived (not duplicate punches)."),
        ("One per record",  "You can request a waiver only once for the same late mark."),
        ("Reason",          "Mandatory — your manager needs context to decide."),
        ("Approval",        "Manager / HR approves or rejects. You'll see the decision in My Requests."),
    ])

    add_heading(doc, "What happens after approval", level=2)
    add_bullet(doc, "The deduction on that late day becomes ₹0 (forgiven).")
    add_bullet(doc, "The record shows a green Approved badge with the approver's name.")

    add_heading(doc, "What happens if rejected", level=2)
    add_bullet(doc, "The original deduction stays in place.")
    add_bullet(doc,
               "You can submit a fresh waiver request for the same record if you have new "
               "context or stronger justification.")

    add_heading(doc, "Common problems", level=2)
    add_bullet(doc,
               "\"No late records found\" — you have no eligible lates in the last 30 days. Nothing to waive.")
    add_bullet(doc,
               "\"The record is greyed out\" — it's already waived or already in an approved waiver.")
    add_bullet(doc,
               "\"Reason field won't accept text\" — make sure you're typing in the multi-line box, not the date field.")

    page_break(doc)

    # ---------- SECTION 5: FIELD ATTENDANCE ----------
    add_heading(doc, "5. Field Attendance", level=1)
    add_para(doc,
             "Use this if your work day starts at a customer site or out on the road — "
             "for example, sales executives, service engineers, or delivery staff who "
             "don't come to the office first.",
             italic=True, colour=GREY)

    add_heading(doc, "Before you can use it", level=2)
    add_bullet(doc,
               "A vehicle trip for today must be assigned to you. Usually your driver "
               "or admin creates this. If you don't see any trip, ask them.")
    add_bullet(doc,
               "You must log at least one customer visit during the day, and GPS must "
               "be turned on so the visit can be tagged with location.")

    add_heading(doc, "Step by step", level=2)
    add_numbered(doc, "Open the app and tap Attendance.")
    add_numbered(doc, "Pick Field Attendance (Customer Visit) and verify yourself.")
    add_numbered(doc, "You'll land on the Today tab.")
    add_numbered(doc, "Tap Mark Field Attendance. The camera opens — take your photo.")
    add_numbered(doc,
                 "Select the Primary Trip: tap Edit, pick the trip assigned to you "
                 "(vehicle + driver + source + destination), and save.")
    add_numbered(doc,
                 "As you visit customers, log each visit through the visits sheet. "
                 "Each visit captures your GPS automatically.")
    add_numbered(doc,
                 "If your route changes during the day, tap Add Another Trip. The "
                 "previous trip closes automatically and any open visits get marked done.")
    add_numbered(doc,
                 "At the end of the day, tap Check Out. Enter the End KM (your "
                 "vehicle's odometer reading) and take a photo. You're done.")

    add_heading(doc, "Rules at a glance", level=2)
    add_rules_table(doc, [
        ("Location",      "GPS is captured for each customer visit, but the 100 m office check does NOT apply here."),
        ("End KM",        "Required before adding a new trip or before checking out at end of day."),
        ("Multiple trips","Each trip can hold multiple visits. You can have several trips in one day."),
        ("Photo",         "Required at check-in and check-out, just like office attendance."),
        ("Late",          "If you check in late, you'll be asked for a reason (same as office attendance)."),
    ])

    add_heading(doc, "The History tab", level=2)
    add_para(doc,
             "Switch to History to see your past field attendance days. You can filter the list:")
    add_bullet(doc, "Pick a From Date and To Date to look at a specific period.")
    add_bullet(doc, "Late Only — show only days you were late.")
    add_bullet(doc, "With Deduction — show only days with a salary deduction.")
    add_bullet(doc, "Waived — show only late days where a waiver was approved.")
    add_para(doc, "Tap Apply to use the filters or Reset to clear them.", italic=True)

    add_heading(doc, "Common problems", level=2)
    add_bullet(doc,
               "\"No trips showing in the picker\" — no trip is assigned to you for today. Ask your driver / admin.")
    add_bullet(doc,
               "\"I forgot to enter End KM\" — the app will prompt you for it automatically before check-out.")
    add_bullet(doc,
               "\"I added a wrong trip\" — tap the delete icon on the trip card and re-add the correct one.")
    add_bullet(doc,
               "\"GPS not working\" — turn on Location in your phone settings; the app needs it for visits.")

    page_break(doc)

    # ---------- APPENDIX ----------
    add_heading(doc, "Appendix — Status colours", level=1)
    add_para(doc,
             "Across Leave Requests, Late Waivers, and other approval flows, the app uses "
             "the same colour code on every status badge:")
    add_status_table(doc)

    add_para(doc,
             "If you ever see a status you don't recognise, tap the request card to open "
             "the full details — there's usually a description in plain text alongside the badge.",
             italic=True, colour=GREY)

    doc.add_paragraph()
    add_para(doc, "— End of manual —", italic=True, colour=GREY,
             size=10)
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
