# -*- coding: utf-8 -*-
"""
Build the three HR module manuals in the "369ai HR Step-by-Step Guide" style,
matching Employee_Monthly_Report_User_Manual.(docx/pdf).

Usage:
    py _build_hr_guides.py --module field
    py _build_hr_guides.py --module late
    py _build_hr_guides.py --module leave
    py _build_hr_guides.py --module all

Text-only (no screenshots). Requires python-docx.
"""
import argparse
import os

from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- design tokens (extracted from the reference) -------------------------
DARK_BLUE = "1F4D79"
MED_BLUE = "2D75B6"
BODY = "333333"
GRAY = "808080"
CAP_GRAY = "7A7A7A"
TIP_FILL, TIP_BORDER = "E2EFDA", "70AD47"
IMP_FILL, IMP_BORDER = "FFF2CC", "E0A800"
NOTE_FILL, NOTE_BORDER = "F2F2F2", "BFBFBF"
WHITE = "FFFFFF"
TBL_BORDER = "BFBFBF"

HERE = os.path.dirname(os.path.abspath(__file__))


# ---- low-level oxml helpers ----------------------------------------------
def _set(el, tag, **attrs):
    child = OxmlElement(tag)
    for k, v in attrs.items():
        child.set(qn(k), str(v))
    el.append(child)
    return child


def cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    _set(tcPr, "w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": fill})


def cell_left_border(cell, color, sz=24):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    _set(borders, "w:left", **{"w:val": "single", "w:sz": sz, "w:space": "0", "w:color": color})
    tcPr.append(borders)


def cell_valign(cell, val="center"):
    tcPr = cell._tc.get_or_add_tcPr()
    _set(tcPr, "w:vAlign", **{"w:val": val})


def table_borders(tbl, color=TBL_BORDER, sz=6):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _set(borders, "w:" + edge, **{"w:val": "single", "w:sz": sz, "w:space": "0", "w:color": color})
    tblPr.append(borders)


def para_bottom_border(p, color, sz=8, space=4):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    _set(pBdr, "w:bottom", **{"w:val": "single", "w:sz": sz, "w:space": space, "w:color": color})
    pPr.append(pBdr)


def para_top_bottom_border(p, color, sz=18, space=8):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    _set(pBdr, "w:top", **{"w:val": "single", "w:sz": sz, "w:space": space, "w:color": color})
    _set(pBdr, "w:bottom", **{"w:val": "single", "w:sz": sz, "w:space": space, "w:color": color})
    pPr.append(pBdr)


def run_shading(run, fill):
    rPr = run._r.get_or_add_rPr()
    _set(rPr, "w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": fill})


def add_run(p, text, *, bold=False, italic=False, color=BODY, size=10.5, shade=None, spacing=None):
    r = p.add_run(text)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if size:
        r.font.size = Pt(size)
    if spacing is not None:
        rPr = r._r.get_or_add_rPr()
        _set(rPr, "w:spacing", **{"w:val": spacing})
    if shade:
        run_shading(r, shade)
    return r


def set_cell_width(cell, twips):
    tcPr = cell._tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn("w:tcW")):
        tcPr.remove(ex)
    _set(tcPr, "w:tcW", **{"w:w": twips, "w:type": "dxa"})


# ---- the guide builder ----------------------------------------------------
class Guide:
    def __init__(self, module_name, subtitle, description, version):
        self.module_name = module_name
        self.doc = Document()
        self._base_style()
        self._page_setup()
        self._header_footer()
        self.content_w = self._content_width_twips()
        self._cover(subtitle, description, version)

    # -- setup --
    def _base_style(self):
        st = self.doc.styles["Normal"]
        st.font.name = "Calibri"
        st.font.size = Pt(10.5)
        st.font.color.rgb = RGBColor.from_string(BODY)
        st.paragraph_format.space_after = Pt(6)

    def _page_setup(self):
        for s in self.doc.sections:
            s.top_margin = Pt(54)
            s.bottom_margin = Pt(54)
            s.left_margin = Pt(57)
            s.right_margin = Pt(57)

    def _content_width_twips(self):
        s = self.doc.sections[0]
        return int((s.page_width - s.left_margin - s.right_margin) / 635)  # EMU->twips

    def _header_footer(self):
        sec = self.doc.sections[0]
        # header
        hdr = sec.header
        hdr.is_linked_to_previous = False
        htbl = hdr.add_table(rows=1, cols=2, width=Twips(self._content_width_twips()))
        htbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        half = self._content_width_twips() // 2
        lc, rc = htbl.cell(0, 0), htbl.cell(0, 1)
        for c in (lc, rc):
            set_cell_width(c, half)
            tcPr = c._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            _set(borders, "w:bottom", **{"w:val": "single", "w:sz": 12, "w:space": "0", "w:color": MED_BLUE})
            tcPr.append(borders)
        lp = lc.paragraphs[0]
        add_run(lp, self.module_name, bold=True, color=DARK_BLUE, size=8.5)
        add_run(lp, "   |   HR Step-by-Step Guide", color=GRAY, size=8.5)
        rp = rc.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(rp, "369ai.Biz", bold=True, color=MED_BLUE, size=8.5)
        # footer
        ftr = sec.footer
        ftr.is_linked_to_previous = False
        fp = ftr.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(fp, "369ai.Biz  ·  %s — HR Guide" % self.module_name, color=GRAY, size=8)

    def _spacer(self, n=1):
        for _ in range(n):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)

    def _cover(self, subtitle, description, version):
        self._spacer(3)
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, "369ai", bold=True, color=MED_BLUE, size=20)
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, "H R   S U I T E   F O R   O D O O   1 9", bold=True, color=GRAY, size=11)
        self._spacer(1)
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_top_bottom_border(p, MED_BLUE)
        add_run(p, self.module_name.upper(), bold=True, color=DARK_BLUE, size=30)
        self._spacer(1)
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, "HR Step-by-Step Guide", color=MED_BLUE, size=18)
        self._spacer(1)
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, subtitle, italic=True, color=GRAY, size=12)
        self._spacer(3)
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, description, italic=True, color=GRAY, size=11)
        self._spacer(1)
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, version, color=GRAY, size=10)
        self.doc.add_page_break()

    # -- components --
    def section_heading(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        para_bottom_border(p, MED_BLUE)
        add_run(p, text, bold=True, color=MED_BLUE, size=15)

    def part(self, text):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_width(cell, self.content_w)
        cell_shading(cell, DARK_BLUE)
        cell_valign(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Pt(6)
        add_run(p, text, bold=True, color=WHITE, size=14)
        self._spacer(1)

    def step(self, n, title):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, "  Step %s  " % n, bold=True, color=WHITE, size=10, shade=MED_BLUE)
        add_run(p, "  ", size=10)
        add_run(p, title, bold=True, color=DARK_BLUE, size=12.5)

    def body(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_run(p, text, color=BODY)
        return p

    def bullets(self, items):
        for it in items:
            p = self.doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            if isinstance(it, tuple):
                add_run(p, it[0], bold=True, color=BODY)
                add_run(p, it[1], color=BODY)
            else:
                add_run(p, it, color=BODY)

    def callout(self, kind, label, text):
        fill, border = {"TIP": (TIP_FILL, TIP_BORDER),
                        "IMPORTANT": (IMP_FILL, IMP_BORDER),
                        "NOTE": (NOTE_FILL, NOTE_BORDER)}[kind]
        lblcolor = {"TIP": TIP_BORDER, "IMPORTANT": IMP_BORDER, "NOTE": GRAY}[kind]
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_width(cell, self.content_w)
        cell_shading(cell, fill)
        cell_left_border(cell, border, sz=24)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Pt(6)
        add_run(p, (label or kind) + "  ", bold=True, color=lblcolor)
        add_run(p, text, color=BODY)
        self._spacer(1)

    def table(self, headers, rows, widths=None):
        ncol = len(headers)
        tbl = self.doc.add_table(rows=1, cols=ncol)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_borders(tbl)
        if widths:
            total = sum(widths)
            widths = [int(self.content_w * w / total) for w in widths]
        for j, h in enumerate(headers):
            c = tbl.cell(0, j)
            if widths:
                set_cell_width(c, widths[j])
            cell_shading(c, DARK_BLUE)
            cell_valign(c)
            pr = c.paragraphs[0]
            pr.paragraph_format.space_before = Pt(2)
            pr.paragraph_format.space_after = Pt(2)
            add_run(pr, h, bold=True, color=WHITE, size=9.5)
        for row in rows:
            cells = tbl.add_row().cells
            for j, val in enumerate(row):
                if widths:
                    set_cell_width(cells[j], widths[j])
                pr = cells[j].paragraphs[0]
                pr.paragraph_format.space_before = Pt(2)
                pr.paragraph_format.space_after = Pt(2)
                bold = (j == 0)
                add_run(pr, val, bold=bold, color=BODY, size=9.5)
        self._spacer(1)

    def things_to_remember(self, items):
        self.section_heading("Things to Remember")
        self.bullets(items)
        self._spacer(1)
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, "— End of Guide —", italic=True, color=GRAY, size=11)
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, "369ai.Biz   |   %s — HR Guide" % self.module_name, color=GRAY, size=9)

    def save(self, path):
        self.doc.save(path)


# ===========================================================================
#  CONTENT
# ===========================================================================
def build_field():
    g = Guide(
        "HR Field Attendance",
        "Attendance built from vehicle trips and customer visits — with GPS, fuel and trip totals.",
        "A simple guide to running a day in the field: check in, log your trips and visits, and check out.",
        "Module Version 19.0.1.15.4   ·   Odoo 19   ·   Human Resources",
    )

    g.section_heading("About This Module")
    g.body("HR Field Attendance is for people who work on the road rather than at a fixed desk — "
           "drivers, sales and service staff who move between the office and customer sites. Instead of a "
           "single check-in, a whole day's journey is captured on one attendance record: the trips taken, "
           "the customers visited, the distance and fuel used, and the GPS location for each leg.")
    g.body("This guide walks through a field day in the order you live it: check in, set up your first "
           "trip, add more trips and visits as the day goes on, head back home, and check out. Everything "
           "is driven by clearly labelled buttons on the attendance form, so you never have to guess the "
           "next step.")

    g.section_heading("What the Module Gives You at a Glance")
    g.bullets([
        ("A Field source flag", " — field attendance is marked separately from manual office attendance."),
        ("One primary trip plus unlimited additional trips", " — a full day of stops on a single record."),
        ("Customer visits", " — each visit is linked to the trip it happened on."),
        ("GPS capture", " — latitude and longitude from the trip or the first visit."),
        ("Trip Totals", " — total distance, duration, fuel litres and fuel amount across every trip."),
        ("Guided buttons & pop-ups", " — set up the primary trip, add trips, and close the previous trip automatically."),
        ("Mobile app support", " — the same flow is available to the phone app over RPC."),
    ])
    g.body("Who this is for: Field Employees who record their day, and HR Officers / Managers who review "
           "field attendance, GPS, visited stops and the trip totals afterwards.")

    g.part("PART 1   Finding Your Field Attendance")
    g.body("Field attendance lives in two convenient places, so you can reach it from the desktop or the "
           "standalone app icon.")
    g.step(1, "Open the Field Attendance list")
    g.body("From the main menu open Attendances → Field Attendance. You land on My Field Attendance, "
           "which shows only your own records. There is also a Field Attendance app icon in the app "
           "launcher (the nine-dot grid) that opens the same screen.")
    g.bullets([
        ("My Field Attendance", " — every employee sees their own field days here."),
        ("All Field Records", " — HR Officers and Managers see everyone's field attendance."),
    ])
    g.callout("TIP", "TIP", "The list shows the date, employee, check-in and check-out (in office time), "
              "the source trip, the visited stops, the location and the trip totals — a quick way to "
              "scan a day before opening it. Use the Late Only, With Deduction and Waived filters to narrow it down.")

    g.part("PART 2   The Field Attendance Form")
    g.body("Opening a field attendance record shows the standard check-in / check-out header plus a Field "
           "Attendance area built from three parts.")
    g.step(2, "Read the Primary Trip Card")
    g.body("The primary trip card is the main trip of the day — usually Home → Office. It shows the "
           "Source Trip, the Source and Destination Location, KM Travelled and Duration, the GPS Latitude "
           "and Longitude, and a Location name. From the card you can Edit Primary Trip, Open Source Trip, "
           "View Visits or Add Fuel.")
    g.step(3, "Read the Additional Trips & Return Home sections")
    g.body("Each extra trip appears as a card with two columns: Trip Details (source, destination, KM, "
           "duration, GPS) on the left and Visit Details (the customer, date, location and GPS of the visit) "
           "on the right. When you start heading back, the return legs are grouped under a Return Home "
           "section and labelled Visit → Office, Office → Home or Visit → Home (Direct).")
    g.step(4, "Read the Trip Totals")
    g.body("The Trip Totals group adds everything up for the day: Total KM Travelled, Total Duration, "
           "Total Fuel (Litres) and Total Fuel Amount. These are calculated automatically from all the "
           "trips on the record.")

    g.part("PART 3   A Day in the Field")
    g.body("This is the heart of the module — the buttons you press, in order, across a working day.")
    g.step(5, "Check in")
    g.body("Start your day with a check-in (from the app or the form). The attendance is marked as a Field "
           "source and the workflow guidance banner appears to point you to the next action.")
    g.step(6, "Enter a late reason (only if you were late)")
    g.body("If you arrived late, scroll to the Late Tracking section and click Enter Late Reason, then type "
           "why. If you were on time, skip this step.")
    g.step(7, "Set up your primary trip")
    g.body("Click Setup Primary Trip (Home to Office). Pick the trip from the Pick a Trip list, enter the "
           "Start KM, optionally name the Location, and click Save. The trip starts automatically and the "
           "primary card fills in. If you are going straight from home to a customer instead, click Setup "
           "Secondary Trip (Home to Visit) and choose the trip and the visit.")
    g.step(8, "Add additional trips as you move")
    g.body("Each time you drive to a new stop, click Add Additional Trip. If a trip is still open, a Close "
           "Previous Trip pop-up appears first — enter the End KM and click Save & Exit; the open trip is "
           "ended automatically and the next trip's pop-up opens. Pick the trip and the visit, enter the "
           "Start KM, and save.")
    g.callout("NOTE", "NOTE", "Adding a new trip always closes the previous one first, so your distance and "
              "fuel figures stay accurate. Ended trips are locked and hidden from the trip picker.")
    g.step(9, "Head back home")
    g.body("When you are ready to return, click Primary Trip (Via Office or Direct) and choose a route: "
           "Via Office (Visit → Office, then Office → Home) or Direct (Visit → Home). For the "
           "Via Office route, after the first leg saves, the Primary Trip (Office to Home) button appears "
           "for the second leg.")
    g.step(10, "More visits before home (optional)")
    g.body("Need one more visit after starting back? Click Add Additional Trip again — the Return Home "
           "section steps aside and a fresh outbound trip begins. Repeat the cycle as many times as the day "
           "needs.")
    g.step(11, "Check out")
    g.body("At the end of the day click Check Out Now at the top. If a trip is still open you are asked for "
           "the End KM first (Save & Exit). On check-out the last trip is closed, every visit is marked "
           "Done, and the whole record locks read-only.")
    g.callout("IMPORTANT", "IMPORTANT", "Once an attendance is checked out it becomes read-only and the "
              "action buttons disappear. Double-check your trips and End KM before you check out.")

    g.part("PART 4   Reviewing & the Mobile App")
    g.step(12, "Review field attendance")
    g.body("Managers open All Field Records to review any employee's day. Group by Employee, Department, "
           "Trip, Date or Month, and filter by My Records, Late Only, With Deduction or Waived. Open any "
           "record to inspect the trips, visits, GPS and trip totals.")
    g.step(13, "Use the mobile app")
    g.body("The phone app drives the same flow through server (RPC) methods — it fetches today's record, "
           "sets up the primary trip, adds additional and return trips, closes open trips and checks out. "
           "The table below lists the main calls.")
    g.table(
        ["Server method", "What it does"],
        [
            ["get_today_field_attendance", "Fetch (or create) today's field attendance for an employee."],
            ["field_action_setup_primary_trip", "Set the primary trip and auto-start it."],
            ["field_action_create_additional_trip", "Add a secondary / additional trip with its visit."],
            ["field_action_create_return_trip", "Add a return-home leg (via office or direct)."],
            ["field_action_close_previous_trip", "End an open trip with its End KM."],
            ["field_action_check_out", "Close the last trip, mark visits done and check out."],
        ],
        widths=[5, 7],
    )

    g.section_heading("Quick Reference — Buttons")
    g.table(
        ["Button", "When to use it"],
        [
            ["Setup Primary Trip (Home to Office)", "Start the day's main trip from home to the office."],
            ["Setup Secondary Trip (Home to Visit)", "Go straight from home to a customer visit."],
            ["Add Additional Trip", "Record the next stop; closes the open trip first."],
            ["Primary Trip (Via Office or Direct)", "Begin the journey home — via office or direct."],
            ["Primary Trip (Office to Home)", "Second leg of the via-office return."],
            ["Open Source Trip / View Visits", "Open the full trip record or the linked visits."],
            ["Add Fuel", "Log fuel against the current trip."],
            ["Check Out Now", "End the day; closes the last trip and locks the record."],
        ],
        widths=[5, 7],
    )

    g.things_to_remember([
        "Check in first — the form guides you to each next button.",
        "Always enter the End KM when closing a trip so distance and fuel stay correct.",
        "Use Via Office or Direct to match how you actually travel home.",
        "Trip Totals add up KM, duration and fuel automatically across every trip.",
        "Checking out marks all visits Done and locks the record — review before you check out.",
    ])
    g.save(os.path.join(HERE, "HR_Field_Attendance_User_Manual.docx"))
    print("Wrote HR_Field_Attendance_User_Manual.docx")


def build_late():
    g = Guide(
        "Attendance Late Tracking",
        "Grace periods, deduction slabs, waivers, split shifts and half-day Fridays — a fair, configurable late policy.",
        "A simple guide to setting up the late policy, reading late records, approving waivers and running the monthly summary.",
        "Module Version 19.0.4.0.0   ·   Odoo 19   ·   Human Resources",
    )

    g.section_heading("About This Module")
    g.body("Attendance Late Tracking & Deductions extends the standard Odoo Attendance app so that every "
           "check-in is measured against your office hours. When someone arrives after their grace time, "
           "the record is flagged late and — once their monthly free allowance is used up — a deduction "
           "is worked out automatically. Staff can give a reason, and managers can waive a deduction when "
           "the lateness was justified.")
    g.body("This guide follows the policy from setup to payroll: configure office hours and deductions, "
           "review late records, handle waiver requests, and generate the monthly late summary.")

    g.section_heading("What the Module Gives You at a Glance")
    g.bullets([
        ("Configurable office start", " — company-wide or per department, in a chosen office timezone."),
        ("Late threshold (grace minutes)", " — how late counts as late."),
        ("Monthly grace quota", " — a number of free late occurrences per month, per session."),
        ("Two deduction modes", " — fixed slab amounts or automatic hourly-wage calculation."),
        ("Single and split shifts", " — each session is tracked on its own."),
        ("Working-days calendar & public holidays", " — only real working days count."),
        ("Half-day Fridays", " — optional short Fridays on the weeks you choose."),
        ("Late reasons & a waiver workflow", " — manager-approved, with a full audit trail."),
        ("Monthly late summary report", " — grouped and filtered, ready for payroll."),
    ])
    g.body("Who this is for: HR Managers who set the policy and approve waivers, HR Officers who review "
           "records and run the report, and Employees who explain a late arrival or raise a waiver.")

    g.part("PART 1   Finding Late Tracking")
    g.step(1, "Open the Late Tracking menu")
    g.body("Everything lives under Attendances → Late Tracking. From there you reach Late Records, "
           "Monthly Summary, Waiver Requests (or My Waiver Requests for staff), and a Configuration group "
           "with Office Hours & Working Days, Deduction Slabs and Public Holidays.")

    g.part("PART 2   Configuring the Late Policy")
    g.body("Set the policy up once. The central screen is Office Hours & Working Days; two supporting "
           "screens define the slab amounts and the holiday calendar.")
    g.step(2, "Set Office Hours & Working Days")
    g.body("Open Configuration → Office Hours & Working Days. Each record defines the policy for a "
           "company, optionally narrowed to one Department, in a chosen Office Timezone.")
    g.bullets([
        ("Scope", " — Company, optional Department, Office Timezone (e.g. Asia/Muscat)."),
        ("Late Settings", " — Late Threshold (minutes) and Grace Late Times/Month."),
        ("Shift Type", " — Single Shift (Session 1 only) or Split Shift (Session 1 + Session 2)."),
        ("Office Hours", " — the Start and End time for each session."),
        ("Working Days", " — tick the days the company actually works."),
        ("Half-Day Friday", " — enable it, set which Fridays (e.g. 2,4) and the shorter hours."),
    ])
    g.callout("TIP", "TIP", "Use the Recompute Late Records button after changing the policy so existing "
              "records pick up the new office hours, grace and slabs.")
    g.step(3, "Choose a Deduction Mode")
    g.body("Two modes decide how much is deducted once an employee is past their grace quota. Pick the one "
           "that matches your payroll rules.")
    g.table(
        ["Deduction mode", "How the amount is worked out"],
        [
            ["Fixed Amount (Slab-based)", "The late minutes are matched to a row in the Deduction Slabs table and that fixed amount is charged."],
            ["Hourly Wage Calculation", "An hourly rate (wage ÷ working days ÷ daily hours) is charged for each started hour late. Falls back to the slab if the employee has no wage."],
        ],
        widths=[4, 8],
    )
    g.step(4, "Set up Deduction Slabs")
    g.body("Open Configuration → Deduction Slabs. Each row is a minute range — Late From and Late To "
           "(set To = 0 for “and above”) — with a Deduction Amount. These are used directly in "
           "Fixed mode, and as a fallback in Hourly mode.")
    g.step(5, "Maintain Public Holidays")
    g.body("Open Configuration → Public Holidays and add each holiday with a Name and Date. Holidays are "
           "removed from the working-days count, so no one is marked late on a day off.")
    g.callout("NOTE", "NOTE", "The grace quota is counted per session, per month. With a split shift, the "
              "first late arrival in Session 1 and the first in Session 2 each use their own free allowance "
              "before any deduction applies.")

    g.part("PART 3   Working with Late Records")
    g.step(6, "Read a late record")
    g.body("Open Late Tracking → Late Records. Each check-in shows whether it Is Late, the Late Time "
           "(H:MM), the Session, whether it is a Half Day, the Late # in the month, the Deduction Amount "
           "and whether it was Waived. Use the Late Only, With Deduction, Missing Reason and group-by "
           "filters to focus.")
    g.step(7, "Enter a late reason")
    g.body("On a late attendance, click Enter Late Reason in the Late Tracking section and type the "
           "explanation, or type straight into the Late Reason field. The button is hidden once a reason is "
           "filled or the record is waived.")

    g.part("PART 4   The Late Waiver Workflow")
    g.body("A waiver lets a manager cancel a late deduction when the lateness was justified. It moves "
           "through a simple, audited set of states.")
    g.step(8, "Submit a waiver (employee)")
    g.body("On a late record, raise a waiver request, fill the Reason for Waiver and click Submit for "
           "Approval. The request moves from Draft to Pending Approval.")
    g.step(9, "Approve or reject (manager)")
    g.body("A manager opens the pending request and clicks Approve Waiver — the attendance is marked "
           "waived and its deduction becomes zero — or Reject with a note. A rejected request can be "
           "reset to Draft and resubmitted.")
    g.table(
        ["State", "Meaning", "Who acts / button"],
        [
            ["Draft", "Being prepared.", "Employee — Submit for Approval"],
            ["Pending Approval", "Waiting on a manager.", "Manager — Approve Waiver / Reject"],
            ["Approved", "Deduction waived (set to zero).", "—"],
            ["Rejected", "Declined, with a reason.", "Employee — Reset to Draft"],
        ],
        widths=[3, 5, 4],
    )

    g.part("PART 5   The Monthly Late Summary Report")
    g.step(10, "Generate the summary")
    g.body("Open Late Tracking → Monthly Summary. Choose the Month and Year, optionally a Department or "
           "specific Employees, then click Generate Summary.")
    g.step(11, "Read the summary")
    g.body("The report lists one row per employee with Total Late Days, Total Late Time (H:MM) and Total "
           "Deductions. Only late arrivals beyond the grace quota are counted — free late times are not "
           "in the totals — so the figures match what payroll should deduct.")
    g.callout("IMPORTANT", "IMPORTANT", "The summary always counts only the late occurrences after the "
              "monthly grace quota, and never counts waived records, so it never over-deducts.")

    g.section_heading("Quick Reference — Menus & Buttons")
    g.table(
        ["Where", "What it does"],
        [
            ["Office Hours & Working Days", "Define office hours, grace, shift type, working days and half-day Fridays."],
            ["Deduction Slabs", "Set the fixed amounts charged for each late-minute range."],
            ["Public Holidays", "Keep the holiday calendar so days off never count as late."],
            ["Late Records", "Review every check-in, the late time and the deduction."],
            ["Enter Late Reason", "Record why an employee was late."],
            ["Waiver Requests", "Submit, approve or reject a waiver of a late deduction."],
            ["Monthly Summary", "Generate the per-employee late totals for payroll."],
            ["Recompute Late Records", "Reapply the policy after a configuration change."],
        ],
        widths=[5, 7],
    )

    g.things_to_remember([
        "Set Office Hours, grace and deduction mode before anything else — everything builds on them.",
        "Grace is per session, per month; deductions start only after it is used up.",
        "Keep Public Holidays current so no one is flagged late on a day off.",
        "Recompute Late Records after any policy change.",
        "Approved waivers zero the deduction; the monthly summary never counts waived or in-grace records.",
    ])
    g.save(os.path.join(HERE, "Attendance_Late_Tracking_User_Manual.docx"))
    print("Wrote Attendance_Late_Tracking_User_Manual.docx")


def build_leave():
    g = Guide(
        "HR Leave Request",
        "Leave requests with manager approval, a paid-leave quota and automatic unpaid deductions.",
        "A simple guide to setting the leave policy, raising a request, approving it and reading the balance.",
        "Module Version 19.0.1.1.0   ·   Odoo 19   ·   Human Resources",
    )

    g.section_heading("About This Module")
    g.body("HR Leave Request gives every employee a simple way to ask for time off and gives managers a "
           "clear way to approve it. Beyond a yes/no, the module understands your paid-leave policy: it "
           "tracks how much paid leave each person has left, splits a request into paid and unpaid days "
           "when the quota runs out, and works out the salary deduction for the unpaid part automatically.")
    g.body("This guide follows a request end to end: configure the policy once, raise a request, approve "
           "or reject it, and read the balance and reports — with REST endpoints for the mobile app at "
           "the end.")

    g.section_heading("What the Module Gives You at a Glance")
    g.bullets([
        ("Simple requests", " — from date, optional to date, half-day option and a reason."),
        ("Six leave types", " — Sick, Casual, Annual, Personal, Emergency and Other."),
        ("A clear lifecycle", " — Draft, Pending Approval, Approved, Rejected and Cancelled."),
        ("Paid-leave quota", " — configurable per year and per month."),
        ("Automatic paid/unpaid split", " — when the quota is exceeded."),
        ("Automatic deductions", " — salary deduction for unpaid days."),
        ("Per-employee balance", " — used and remaining at a glance."),
        ("Overlap protection", " — two requests cannot cover the same dates."),
        ("Carry-forward & REST API", " — optional carry-over and mobile-app endpoints."),
    ])
    g.body("Who this is for: Employees who raise requests, Managers who approve them, and HR "
           "Administrators who set the company leave policy and quotas.")

    g.part("PART 1   Finding Leave Requests")
    g.step(1, "Open the Leave Requests menu")
    g.body("Everything lives under Attendances → Leave Requests. Employees use My Leave Requests; "
           "managers also see All Leave Requests and the Approved Leaves Report. The policy sits under "
           "Configuration → Leave Policy.")

    g.part("PART 2   Configuring the Leave Policy")
    g.body("The policy is set once per company and drives every calculation that follows.")
    g.step(2, "Set the paid-leave settings")
    g.body("Open Configuration → Leave Policy. Turn on Enable Paid Leave and set Paid Leave Days Per "
           "Year and Paid Leave Days Per Month. These define how much leave is paid before deductions begin.")
    g.step(3, "Set the unpaid / deduction settings")
    g.body("Turn on Enable Unpaid Leave Deduction so that days beyond the quota are deducted. The "
           "deduction uses the employee's monthly wage divided by the working days in the month; a half "
           "day is half of that daily rate.")
    g.step(4, "Set carry-forward (optional)")
    g.body("Enable Allow Carry Forward and set Max Carry Forward Days if unused paid leave should roll "
           "over to next year.")
    g.callout("NOTE", "NOTE", "Paid days are limited by the smaller of the yearly and monthly remaining "
              "quota. Anything above that becomes unpaid and is deducted automatically.")

    g.part("PART 3   Raising a Leave Request")
    g.body("This part is for employees. Open Attendances → Leave Requests → My Leave Requests.")
    g.step(5, "Fill in the request")
    g.body("Click New and complete the form: confirm your Employee, choose a Leave Type (Casual is the "
           "default), set the From Date (required) and, for more than one day, a To Date. Tick Half Day "
           "for a half-day request, and add a Reason.")
    g.table(
        ["Leave type", "Typical use"],
        [
            ["Sick Leave", "Illness or medical appointments."],
            ["Casual Leave", "Short, planned personal time off (the default)."],
            ["Annual Leave", "Planned holiday / vacation."],
            ["Personal Leave", "Personal matters not covered above."],
            ["Emergency Leave", "Urgent, unforeseen situations."],
            ["Other", "Anything that does not fit the categories above."],
        ],
        widths=[4, 8],
    )
    g.body("As you fill the dates the form shows Number of Days, and once saved it splits them into Paid "
           "Days and Unpaid Days with the matching Deduction Amount.")
    g.step(6, "Submit for approval")
    g.body("Click Submit for Approval. The request moves from Draft to Pending Approval and waits for a "
           "manager. You can Cancel a request, and a rejected or cancelled one can be reset to Draft and "
           "resubmitted.")
    g.callout("IMPORTANT", "IMPORTANT", "Requests cannot overlap. If your dates clash with an existing "
              "non-rejected request, the system blocks it and tells you which request it conflicts with.")

    g.part("PART 4   Approving Requests")
    g.step(7, "Review and decide (manager)")
    g.body("Open All Leave Requests — it opens pre-filtered to Pending. Open a request, check the dates, "
           "type, paid/unpaid split and reason, then click Approve or Reject. Rejecting lets you add a "
           "Rejection Reason that the employee will see.")
    g.table(
        ["State", "Meaning", "Who acts / button"],
        [
            ["Draft", "Being prepared.", "Employee — Submit for Approval"],
            ["Pending Approval", "Waiting on a manager.", "Manager — Approve / Reject"],
            ["Approved", "Granted; paid/unpaid split fixed.", "Employee — Cancel (if needed)"],
            ["Rejected", "Declined, with a reason.", "Employee — Reset to Draft"],
            ["Cancelled", "Withdrawn.", "Employee — Reset to Draft"],
        ],
        widths=[3, 5, 4],
    )

    g.part("PART 5   Balance, Reporting & Mobile App")
    g.step(8, "Read the balance and reports")
    g.body("My Leave Requests lists your requests colour-coded by status (green approved, orange pending, "
           "red rejected, grey cancelled) with the paid, unpaid and deduction figures. Managers use the "
           "Approved Leaves Report, grouped by employee, for payroll and the leave calendar.")
    g.step(9, "Use the mobile app")
    g.body("The phone app talks to the module over REST endpoints — create a request, list your "
           "requests, see pending approvals, approve, reject, cancel and pull a report.")
    g.table(
        ["Endpoint", "What it does"],
        [
            ["/leave/request/create", "Submit a new leave request."],
            ["/leave/request/my_requests", "List the signed-in employee's requests."],
            ["/leave/request/pending", "List requests awaiting a manager."],
            ["/leave/request/approve", "Approve a request."],
            ["/leave/request/reject", "Reject a request, with an optional reason."],
            ["/leave/request/cancel", "Cancel a request."],
            ["/leave/request/report", "Pull a filtered leave report."],
        ],
        widths=[5, 7],
    )

    g.section_heading("Quick Reference — Buttons & Menus")
    g.table(
        ["Where", "What it does"],
        [
            ["Leave Policy", "Set paid quota, unpaid deduction and carry-forward (admin)."],
            ["My Leave Requests", "Raise and track your own requests."],
            ["Submit for Approval", "Send a draft request to your manager."],
            ["Approve / Reject", "Manager decision on a pending request."],
            ["Cancel / Reset to Draft", "Withdraw a request, or reopen a rejected/cancelled one."],
            ["All Leave Requests", "Manager queue, pre-filtered to pending."],
            ["Approved Leaves Report", "Approved leave grouped by employee for payroll."],
        ],
        widths=[5, 7],
    )

    g.things_to_remember([
        "Set the Leave Policy first — the paid/unpaid split and deductions depend on it.",
        "From Date is required; add a To Date for multi-day leave, or tick Half Day for half a day.",
        "Paid days follow your remaining yearly and monthly quota; the rest is unpaid and deducted.",
        "Requests cannot overlap an existing non-rejected request.",
        "Managers work from All Leave Requests (pending); Approved Leaves Report feeds payroll.",
    ])
    g.save(os.path.join(HERE, "HR_Leave_Request_User_Manual.docx"))
    print("Wrote HR_Leave_Request_User_Manual.docx")


BUILDERS = {"field": build_field, "late": build_late, "leave": build_leave}

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--module", choices=list(BUILDERS) + ["all"], default="all")
    args = ap.parse_args()
    targets = list(BUILDERS) if args.module == "all" else [args.module]
    for t in targets:
        BUILDERS[t]()
